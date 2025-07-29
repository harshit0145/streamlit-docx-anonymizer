import re
from docx import Document
from faker import Faker
import streamlit as st
from io import BytesIO
from adds import generate_malaysian_address

faker = Faker()

def replace_addresses(text):
    flat_text = text.replace("\n", " ").replace("\r", " ")
    address_matches = re.findall(r"No\.?\s?.{10,120}?\d{5}\s?\w+", flat_text)
    for match in address_matches:
        fake_address = generate_malaysian_address()

        if match in text:
            text = text.replace(match, fake_address)
        elif match in flat_text:
            text = re.sub(re.escape(match), fake_address, text, flags=re.MULTILINE)

    location_replacements = [
        "Kuah Seksyen 9, Tempat Kelibang, Langkawi,Negeri Kedah",
        "Bandar Kuah, Seksyen 9, Tempat Kelibang, Langkawi, Negeri Kedah",
        "Bandar, Seksyen 9, Tempat Kelibang, Daerah Langkawi, Negeri Kedah"
    ]
    for loc in location_replacements:
        fake_address = generate_malaysian_address()
        text = text.replace(loc, fake_address)

    return text

def replace_dates(text, year="2025"):
    text = re.sub(r"\b(\d{1,2})(st|nd|rd|th)? day of (\w+) \d{4}\b", rf"\1 day of \3 {year}", text)
    text = re.sub(r"\b(January|February|March|April|May|June|July|August|September|October|November|December) \d{4}\b", rf"\1 {year}", text)
    return text

def anonymize_docx(file, year="2025"):
    doc = Document(file)


    main_party_name = "COMPANY YYY SDN. BHD."
    counter_party_name = "COMPANY AAA SDN. BHD."
    main_party_number = "998877-U"
    counter_party_number = "112233-A"
    university_name = "Company XYZ "
    replace_name = "private company " 
    

    replacements = {
        "COMPANY ABC SDN. BHD. (COMPANY NO: 880089-U)": f"{main_party_name} (Company No: {main_party_number})",
        "COMPANY XYZ SDN BHD (Company No. 551966-A)": f"{counter_party_name} (Company No: {counter_party_number})",


        "COMPANY ABC": main_party_name,
        "COMPANY XYZ": counter_party_name,
        "Co. No. 101067-P" : main_party_number,


        "Public University ": replace_name,
        "University": "company ",
        "Universities ": "companys  ",
        "Colleges ": " ",
        "ABC " : " XYZ ",
        "UNIVERSITY" : "COMPANY ",
        "ABC University Malaysia": university_name,
    }

    for para in doc.paragraphs:
        text = para.text
        text = replace_addresses(text)
        text = replace_dates(text, year)
        for old, new in replacements.items():
            text = text.replace(old, new)
        para.text = text


    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    text = replace_addresses(text)
                    text = replace_dates(text, year)
                    for old, new in replacements.items():
                        text = text.replace(old, new)
                    para.text = text

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

st.set_page_config(page_title="Legal Agreement Anonymizer ", page_icon="📄", layout="centered")

st.markdown("""
# 📄 Legal Agreement Anonymizer (Malaysia Style)

Welcome to the **Legal Agreement Anonymizer**!  
This tool helps you redact and anonymize legal agreements in a Malaysia-compliant format.

📌 Upload a `.docx` file, and we'll:
- Replace sensitive addresses with Malaysian-style fake ones
- Standardize parties and company numbers
- Neutralize date references (to year 2025)
""")

st.markdown("---")

uploaded_file = st.file_uploader("📤 Upload your `.docx` agreement file below", type=["docx"])

if uploaded_file:
    st.success("✅ File uploaded successfully!")

    if st.button("🔐 Convert Agreement"):
        with st.spinner("Processing and anonymizing the agreement..."):
            output = anonymize_docx(uploaded_file)

        st.success("✅ Agreement Converted successfully!")
        st.download_button(
            label="📥 Download Converted Agreement",
            data=output,
            file_name="Anonymized_Agreement.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
